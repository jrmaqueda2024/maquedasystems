"""
reportes_formatos.py
Generadores de reportes en distintos formatos de archivo, todos
consumiendo la estructura neutral de reportes_datos.py:
- PDF simple (sin gráfico, solo texto/tablas) — reportlab
- Word (.docx) — python-docx
- LibreOffice (.odt) — odfpy
- CSV — csv (librería estándar)
- JSON backup — json (librería estándar)

Requiere: pip install reportlab python-docx odfpy
"""
import csv
import json
import datetime

AZUL_HEX = "1d5fd6"


# ==================== PDF SIMPLE (sin gráfico) ====================
def generar_pdf_simple(ruta_destino: str, datos: dict) -> str:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER

    azul = colors.HexColor(f"#{AZUL_HEX}")
    gris_texto = colors.HexColor("#555555")

    doc = SimpleDocTemplate(
        ruta_destino, pagesize=A4,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm, leftMargin=1.5 * cm, rightMargin=1.5 * cm,
    )
    estilos = getSampleStyleSheet()
    estilo_titulo = ParagraphStyle("Titulo", parent=estilos["Title"], fontSize=16, textColor=azul)
    estilo_subtitulo = ParagraphStyle("Subtitulo", parent=estilos["Normal"], fontSize=9,
                                       textColor=gris_texto, alignment=TA_CENTER)
    estilo_seccion = ParagraphStyle("Seccion", parent=estilos["Heading2"], fontSize=11,
                                     textColor=colors.white, backColor=azul, spaceBefore=10,
                                     spaceAfter=6, leftIndent=4, borderPadding=4)
    estilo_normal = estilos["Normal"]

    elementos = [Paragraph(datos["titulo"], estilo_titulo), Paragraph(datos["subtitulo"], estilo_subtitulo)]
    if datos.get("generado_por"):
        elementos.append(Paragraph(f"Generado por {datos['generado_por']}", estilo_subtitulo))
    elementos.append(Spacer(1, 12))

    for seccion in datos["secciones"]:
        elementos.append(Paragraph(seccion["titulo"], estilo_seccion))
        if seccion["tipo"] == "resumen":
            filas = [[etiqueta, valor] for etiqueta, valor in seccion["filas"]]
            t = Table(filas, colWidths=[10 * cm, 6 * cm])
            t.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (1, 0), (1, -1), "RIGHT"), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LINEBELOW", (0, 0), (-1, -2), 0.4, colors.HexColor("#dddddd")),
            ]))
            elementos.append(t)
        elif seccion["tipo"] == "grafico_barras":
            elementos.append(Paragraph("(Gráfico omitido en el PDF simple; ver PDF con dashboard)", estilo_normal))
        elif seccion["tipo"] == "tabla":
            if seccion["filas"]:
                filas = [seccion["encabezados"]] + seccion["filas"]
                t = Table(filas, repeatRows=1)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), azul), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f4f5f7")]),
                    ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#dddddd")),
                    ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                elementos.append(t)
            else:
                elementos.append(Paragraph("Sin datos.", estilo_normal))
        elementos.append(Spacer(1, 8))

    doc.build(elementos)
    return ruta_destino


# ==================== WORD (.docx) ====================
def generar_word(ruta_destino: str, datos: dict) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    color_azul = RGBColor(0x1D, 0x5F, 0xD6)

    doc = Document()
    for seccion_estilo in doc.styles:
        if seccion_estilo.name == "Normal":
            seccion_estilo.font.name = "Arial"
            seccion_estilo.font.size = Pt(10)

    titulo = doc.add_heading(datos["titulo"], level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = color_azul

    p_sub = doc.add_paragraph(datos["subtitulo"])
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.runs[0].font.size = Pt(9)
    p_sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if datos.get("generado_por"):
        p_gen = doc.add_paragraph(f"Generado por {datos['generado_por']}")
        p_gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_gen.runs[0].font.size = Pt(9)
        p_gen.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for seccion in datos["secciones"]:
        encabezado = doc.add_heading(seccion["titulo"], level=1)
        for run in encabezado.runs:
            run.font.color.rgb = color_azul
            run.font.size = Pt(13)

        if seccion["tipo"] == "resumen":
            tabla = doc.add_table(rows=0, cols=2)
            tabla.style = "Light Grid Accent 1"
            for etiqueta, valor in seccion["filas"]:
                fila = tabla.add_row().cells
                fila[0].text = etiqueta
                fila[1].text = valor
                fila[0].paragraphs[0].runs[0].font.bold = True
        elif seccion["tipo"] == "grafico_barras":
            doc.add_paragraph("(Gráfico omitido en este formato; ver PDF con dashboard o Excel con gráficos)")
        elif seccion["tipo"] == "tabla":
            if seccion["filas"]:
                tabla = doc.add_table(rows=1, cols=len(seccion["encabezados"]))
                tabla.style = "Light Grid Accent 1"
                for i, texto in enumerate(seccion["encabezados"]):
                    celda = tabla.rows[0].cells[i]
                    celda.text = texto
                    celda.paragraphs[0].runs[0].font.bold = True
                for fila_datos in seccion["filas"]:
                    fila = tabla.add_row().cells
                    for i, valor in enumerate(fila_datos):
                        fila[i].text = str(valor)
            else:
                doc.add_paragraph("Sin datos.")
        doc.add_paragraph()

    doc.save(ruta_destino)
    return ruta_destino


# ==================== LIBREOFFICE (.odt) ====================
def generar_odt(ruta_destino: str, datos: dict) -> str:
    from odf.opendocument import OpenDocumentText
    from odf.text import H, P
    from odf.table import Table, TableRow, TableCell
    from odf.style import (
        Style, TextProperties, ParagraphProperties,
        TableCellProperties,
    )

    doc = OpenDocumentText()

    estilo_titulo = Style(name="TituloReporte", family="paragraph")
    estilo_titulo.addElement(TextProperties(fontsize="18pt", fontweight="bold", color=f"#{AZUL_HEX}"))
    doc.styles.addElement(estilo_titulo)

    estilo_subtitulo = Style(name="SubtituloReporte", family="paragraph")
    estilo_subtitulo.addElement(TextProperties(fontsize="9pt", color="#555555"))
    doc.styles.addElement(estilo_subtitulo)

    estilo_seccion = Style(name="SeccionReporte", family="paragraph")
    estilo_seccion.addElement(TextProperties(fontsize="12pt", fontweight="bold", color="#ffffff"))
    estilo_seccion.addElement(ParagraphProperties(backgroundcolor=f"#{AZUL_HEX}"))
    doc.styles.addElement(estilo_seccion)

    estilo_celda_encabezado = Style(name="CeldaEncabezado", family="table-cell")
    estilo_celda_encabezado.addElement(TableCellProperties(backgroundcolor=f"#{AZUL_HEX}"))
    doc.styles.addElement(estilo_celda_encabezado)

    doc.text.addElement(H(outlinelevel=1, stylename=estilo_titulo, text=datos["titulo"]))
    doc.text.addElement(P(stylename=estilo_subtitulo, text=datos["subtitulo"]))
    if datos.get("generado_por"):
        doc.text.addElement(P(stylename=estilo_subtitulo, text=f"Generado por {datos['generado_por']}"))

    for seccion in datos["secciones"]:
        doc.text.addElement(P(stylename=estilo_seccion, text=seccion["titulo"]))

        if seccion["tipo"] == "resumen":
            tabla = Table(name=f"tabla_{id(seccion)}")
            for etiqueta, valor in seccion["filas"]:
                fila = TableRow()
                celda1 = TableCell()
                celda1.addElement(P(text=etiqueta))
                celda2 = TableCell()
                celda2.addElement(P(text=valor))
                fila.addElement(celda1)
                fila.addElement(celda2)
                tabla.addElement(fila)
            doc.text.addElement(tabla)
        elif seccion["tipo"] == "grafico_barras":
            doc.text.addElement(P(text="(Gráfico omitido en este formato; ver PDF con dashboard o Excel con gráficos)"))
        elif seccion["tipo"] == "tabla":
            if seccion["filas"]:
                tabla = Table(name=f"tabla_{id(seccion)}")
                fila_encabezado = TableRow()
                for texto in seccion["encabezados"]:
                    celda = TableCell(stylename=estilo_celda_encabezado)
                    celda.addElement(P(text=texto))
                    fila_encabezado.addElement(celda)
                tabla.addElement(fila_encabezado)
                for fila_datos in seccion["filas"]:
                    fila = TableRow()
                    for valor in fila_datos:
                        celda = TableCell()
                        celda.addElement(P(text=str(valor)))
                        fila.addElement(celda)
                    tabla.addElement(fila)
                doc.text.addElement(tabla)
            else:
                doc.text.addElement(P(text="Sin datos."))
        doc.text.addElement(P(text=""))

    doc.save(ruta_destino)
    return ruta_destino


# ==================== CSV ====================
def generar_csv(ruta_destino: str, datos: dict) -> str:
    """Genera un único CSV con todas las secciones concatenadas, separadas
    por una fila de título en mayúsculas (formato simple, legible en Excel
    o cualquier hoja de cálculo)."""
    with open(ruta_destino, "w", newline="", encoding="utf-8-sig") as f:
        escritor = csv.writer(f)
        escritor.writerow([datos["titulo"]])
        escritor.writerow([datos["subtitulo"]])
        if datos.get("generado_por"):
            escritor.writerow([f"Generado por {datos['generado_por']}"])
        escritor.writerow([])

        for seccion in datos["secciones"]:
            escritor.writerow([seccion["titulo"]])
            if seccion["tipo"] == "resumen":
                for etiqueta, valor in seccion["filas"]:
                    escritor.writerow([etiqueta, valor])
            elif seccion["tipo"] == "grafico_barras":
                escritor.writerow(["Fecha"] + seccion["categorias"])
                escritor.writerow(["Total"] + [f"{v:.0f}" for v in seccion["valores"]])
            elif seccion["tipo"] == "tabla":
                escritor.writerow(seccion["encabezados"])
                for fila in seccion["filas"]:
                    escritor.writerow(fila)
            escritor.writerow([])

    return ruta_destino


# ==================== JSON (backup) ====================
def generar_json(ruta_destino: str, datos: dict) -> str:
    """Genera un backup estructurado en JSON con todos los datos del
    reporte, útil para integraciones externas o respaldo de información."""
    backup = {
        "titulo": datos["titulo"],
        "subtitulo": datos["subtitulo"],
        "generado_por": datos.get("generado_por", ""),
        "fecha_exportacion": datetime.datetime.now().isoformat(),
        "secciones": datos["secciones"],
    }
    with open(ruta_destino, "w", encoding="utf-8") as f:
        json.dump(backup, f, ensure_ascii=False, indent=2)
    return ruta_destino